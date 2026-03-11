import requests
import os
import re
import docx

# Configuration
OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "qwen2.5-coder:7b" 
OUTPUT_DIR = "pending_solutions"
DOCX_PATH = r"C:\Users\hp\Documents\Life\questions.docx"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# 1. Improved Extraction Engine (Captures almost any line with a question)
questions = []
print(f"Reading: {DOCX_PATH}")

try:
    doc = docx.Document(DOCX_PATH)
    for para in doc.paragraphs:
        line = para.text.strip()
        
        # This will capture "1. Question", "Q1. Question", or even just "Question" 
        # as long as it isn't a section header like "1-10: Basic"
        if line and not line.endswith(':') and not "Section" in line:
            # Clean out common Word artifacts
            clean_q = re.sub(r'^\d+[\.\s\-]+|Q\d+[\.\s\-]+', '', line).strip()
            if clean_q:
                questions.append(clean_q)
except Exception as e:
    print(f"Error: {e}")
    exit()

# Safety check: If it still finds 0, it takes every non-empty line
if len(questions) == 0:
    questions = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

print(f"Extracted {len(questions)} items. Starting high-speed generation...")

# 2. Optimized Generation Loop
for i, question in enumerate(questions):
    day_num = i + 1
    # We only want to generate up to 300
    if day_num > 300: break
    
    filename = f"day_{day_num:03d}.py"
    filepath = os.path.join(OUTPUT_DIR, filename)
    
    if os.path.exists(filepath):
        continue

    print(f"Generating {filename}...")
    
    # Strict prompt for clean, professional code
    prompt = f"""
    Question: {question}
    
    Task: Write a perfect Python solution.
    Rules:
    1. Output ONLY the Python code. No 'Here is the code' or 'Explanation'.
    2. Start with the question text in a docstring: \"\"\" {question} \"\"\"
    3. Use 'https://example.com' for any URLs.
    4. Remove all conversational comments. 
    5. Ensure code is professional and error-free.
    """

    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "stream": False
    }

    try:
        response = requests.post(OLLAMA_URL, json=payload, timeout=60)
        response.raise_for_status()
        
        raw_output = response.json().get("response", "")
        # Remove any markdown triple backticks the AI might add
        clean_code = re.sub(r'```python|```', '', raw_output).strip()
        
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(clean_code)
            
    except Exception as e:
        print(f"Failed at {filename}: {e}")
        break

print(f"\nSuccess! Files are in {OUTPUT_DIR}")